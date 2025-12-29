"""
Query Package Builder
=====================

PHASE 2 FEATURE

Generates agent submission packages:
- Query letter (DOCX with template)
- Synopsis (1-2 pages)
- Sample chapters (first 3 or 50 pages)
- Comp titles with reasoning
- Author bio

Output: exports/query_packages/
"""

import json
import shutil
from pathlib import Path
from datetime import datetime
from typing import Dict, List, Optional
from dataclasses import dataclass, asdict


@dataclass
class QueryPackageConfig:
    """Configuration for query package."""
    title: str = "Under the Big Top"
    author_name: str = ""
    author_email: str = ""
    word_count: int = 80000
    genre: str = "Literary Historical Fiction"
    sample_chapters: List[int] = None  # Which chapters to include
    comp_titles: List[str] = None
    
    def __post_init__(self):
        if self.sample_chapters is None:
            self.sample_chapters = [1, 2, 3]
        if self.comp_titles is None:
            self.comp_titles = [
                "Water for Elephants by Sara Gruen",
                "The Night Circus by Erin Morgenstern"
            ]


@dataclass  
class QueryPackage:
    """Generated query package."""
    package_id: str
    created_at: str
    config: QueryPackageConfig
    files: Dict[str, str]  # filename -> path


class QueryBuilder:
    """
    Builds agent submission packages.
    
    Components:
    1. Query letter (personalized template)
    2. Synopsis (AI-generated from chapter summaries)
    3. Sample chapters (formatted to spec)
    4. Comp titles document
    5. Author bio
    """
    
    QUERY_TEMPLATE = """Dear [AGENT NAME],

I am seeking representation for {title}, a {genre} complete at {word_count:,} words.

[HOOK PARAGRAPH - 1-2 sentences that grab attention]

[MAIN PLOT PARAGRAPH - What's at stake, who's the protagonist, what do they want]

{title} will appeal to readers who enjoyed {comps}.

[AUTHOR BIO - Relevant credentials, why you wrote this]

Thank you for your time and consideration. I have included [MATERIALS] per your submission guidelines.

Sincerely,
{author_name}
{author_email}
"""

    SYNOPSIS_TEMPLATE = """SYNOPSIS: {title}
by {author_name}
{genre} | {word_count:,} words

---

[OPENING - Introduce protagonist and world]

[INCITING INCIDENT - What disrupts the status quo]

[RISING ACTION - Major plot points, 2-3 paragraphs]

[CLIMAX - The breaking point]

[RESOLUTION - How it ends, what changes]

---

{title} explores themes of {themes}.
"""
    
    def __init__(self, base_dir: Path):
        self.base_dir = Path(base_dir)
        self.chapters_dir = self.base_dir / "chapters"
        self.exports_dir = self.base_dir / "exports"
        self.packages_dir = self.exports_dir / "query_packages"
        self.packages_dir.mkdir(parents=True, exist_ok=True)
    
    def _generate_package_id(self) -> str:
        """Generate unique package ID."""
        return datetime.now().strftime("query_%Y%m%d_%H%M%S")
    
    def _create_query_letter(self, config: QueryPackageConfig, package_dir: Path) -> str:
        """Create query letter from template."""
        comps = " and ".join(config.comp_titles[:2])
        
        content = self.QUERY_TEMPLATE.format(
            title=config.title,
            genre=config.genre,
            word_count=config.word_count,
            comps=comps,
            author_name=config.author_name or "[YOUR NAME]",
            author_email=config.author_email or "[YOUR EMAIL]"
        )
        
        # Save as DOCX
        try:
            from docx import Document
            from docx.shared import Pt
            
            doc = Document()
            for para_text in content.split('\n\n'):
                p = doc.add_paragraph(para_text.strip())
                for run in p.runs:
                    run.font.size = Pt(12)
                    run.font.name = 'Times New Roman'
            
            filepath = package_dir / "query_letter.docx"
            doc.save(filepath)
            return str(filepath)
        except ImportError:
            # Fallback to txt
            filepath = package_dir / "query_letter.txt"
            filepath.write_text(content)
            return str(filepath)
    
    def _create_synopsis(self, config: QueryPackageConfig, package_dir: Path) -> str:
        """Create synopsis template."""
        content = self.SYNOPSIS_TEMPLATE.format(
            title=config.title,
            author_name=config.author_name or "[YOUR NAME]",
            genre=config.genre,
            word_count=config.word_count,
            themes="silence as complicity, jealousy, and dignity in labor"
        )
        
        try:
            from docx import Document
            from docx.shared import Pt
            
            doc = Document()
            for para_text in content.split('\n\n'):
                p = doc.add_paragraph(para_text.strip())
            
            filepath = package_dir / "synopsis.docx"
            doc.save(filepath)
            return str(filepath)
        except ImportError:
            filepath = package_dir / "synopsis.txt"
            filepath.write_text(content)
            return str(filepath)
    
    def _create_sample_chapters(self, config: QueryPackageConfig, package_dir: Path) -> str:
        """Compile sample chapters into single document."""
        try:
            from docx import Document
            from docx.shared import Pt, Inches
            from docx.enum.text import WD_ALIGN_PARAGRAPH
            
            doc = Document()
            
            # Title page
            title_para = doc.add_paragraph()
            title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = title_para.add_run(config.title.upper())
            run.bold = True
            run.font.size = Pt(18)
            
            doc.add_paragraph()
            author_para = doc.add_paragraph()
            author_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            author_para.add_run(f"by {config.author_name or '[YOUR NAME]'}")
            
            doc.add_page_break()
            
            # Add chapters
            for ch_num in config.sample_chapters:
                # Read chapter
                chapter_text = None
                for ext in ['.md', '.docx']:
                    path = self.chapters_dir / f"chapter_{ch_num:02d}{ext}"
                    if path.exists():
                        if ext == '.md':
                            chapter_text = path.read_text(encoding='utf-8')
                        else:
                            try:
                                ch_doc = Document(path)
                                chapter_text = '\n\n'.join([p.text for p in ch_doc.paragraphs])
                            except:
                                pass
                        break
                
                if chapter_text:
                    # Add chapter heading
                    heading = doc.add_heading(f"Chapter {ch_num}", level=1)
                    heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    
                    # Add content
                    for para_text in chapter_text.split('\n\n'):
                        para_text = para_text.strip()
                        if para_text and not para_text.startswith('#'):
                            p = doc.add_paragraph(para_text)
                            p.paragraph_format.first_line_indent = Inches(0.5)
                    
                    doc.add_page_break()
            
            filepath = package_dir / "sample_chapters.docx"
            doc.save(filepath)
            return str(filepath)
            
        except ImportError:
            return ""
    
    def _create_comp_titles(self, config: QueryPackageConfig, package_dir: Path) -> str:
        """Create comp titles document."""
        content = [
            f"COMPARABLE TITLES for {config.title}",
            "=" * 50,
            ""
        ]
        
        comp_notes = {
            "Water for Elephants by Sara Gruen": 
                "Similar setting (traveling circus), historical period (Depression era vs 1950s), "
                "love triangle structure, themes of labor and belonging.",
            "The Night Circus by Erin Morgenstern":
                "Circus atmosphere, romantic tension, literary prose style, "
                "emphasis on wonder and danger of performance."
        }
        
        for comp in config.comp_titles:
            content.append(f"• {comp}")
            if comp in comp_notes:
                content.append(f"  {comp_notes[comp]}")
            content.append("")
        
        filepath = package_dir / "comp_titles.txt"
        filepath.write_text('\n'.join(content))
        return str(filepath)
    
    def _create_bio(self, config: QueryPackageConfig, package_dir: Path) -> str:
        """Create author bio template."""
        content = f"""AUTHOR BIO

{config.author_name or '[YOUR NAME]'}

[Write 2-3 sentences about yourself. Include:]
- Relevant background or expertise (if any)
- Why you wrote this particular story
- Previous publications (if any)
- Location (optional)

Example:
"{config.author_name or 'Author'} has been fascinated by circus history since 
discovering a collection of 1950s circus programs in their grandmother's attic. 
This novel grew from that discovery and years of research into the lives of 
circus workers during the Bracero era. They live in [CITY] with [DETAILS]."
"""
        filepath = package_dir / "author_bio.txt"
        filepath.write_text(content)
        return str(filepath)
    
    def build_package(self, config: QueryPackageConfig) -> QueryPackage:
        """Build complete query package."""
        package_id = self._generate_package_id()
        package_dir = self.packages_dir / package_id
        package_dir.mkdir(exist_ok=True)
        
        files = {}
        
        # Generate each component
        files["query_letter"] = self._create_query_letter(config, package_dir)
        files["synopsis"] = self._create_synopsis(config, package_dir)
        files["sample_chapters"] = self._create_sample_chapters(config, package_dir)
        files["comp_titles"] = self._create_comp_titles(config, package_dir)
        files["author_bio"] = self._create_bio(config, package_dir)
        
        # Save config
        config_path = package_dir / "package_config.json"
        config_path.write_text(json.dumps(asdict(config), indent=2))
        
        package = QueryPackage(
            package_id=package_id,
            created_at=datetime.now().isoformat(),
            config=config,
            files=files
        )
        
        return package
    
    def list_packages(self) -> List[str]:
        """List existing query packages."""
        return [d.name for d in self.packages_dir.iterdir() if d.is_dir()]


# Convenience functions

def build_query_package(
    author_name: str = "",
    author_email: str = "",
    sample_chapters: List[int] = None,
    base_dir: str = "."
) -> Dict:
    """Build a query package."""
    config = QueryPackageConfig(
        author_name=author_name,
        author_email=author_email,
        sample_chapters=sample_chapters or [1, 2, 3]
    )
    
    builder = QueryBuilder(Path(base_dir))
    package = builder.build_package(config)
    
    return {
        "package_id": package.package_id,
        "created_at": package.created_at,
        "files": package.files,
        "location": str(builder.packages_dir / package.package_id)
    }


def list_query_packages(base_dir: str = ".") -> List[str]:
    """List existing query packages."""
    builder = QueryBuilder(Path(base_dir))
    return builder.list_packages()
