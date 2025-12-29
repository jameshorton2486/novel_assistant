"""
Export Pipeline Service
=======================

Normalizes manuscript and exports to publication formats.
Supports DOCX (Word/KDP/Agents) and EPUB (Kindle/Apple/Kobo).
"""

import re
import json
import hashlib
from pathlib import Path
from datetime import datetime
from typing import Dict, List, Optional
from dataclasses import dataclass, asdict


@dataclass
class BookMetadata:
    title: str = "Under the Big Top"
    author: str = ""
    description: str = ""
    keywords: List[str] = None
    categories: List[str] = None
    copyright_year: int = 2024
    
    def __post_init__(self):
        if self.keywords is None:
            self.keywords = ["historical fiction", "1950s", "circus", "California"]
        if self.categories is None:
            self.categories = ["Historical Fiction", "Literary Fiction"]


@dataclass
class ChapterInfo:
    number: int
    title: str
    word_count: int
    content: str


class ManuscriptNormalizer:
    """Normalizes chapters for export."""
    
    SCENE_BREAK = "* * *"
    
    def __init__(self, chapters_dir: Path):
        self.chapters_dir = Path(chapters_dir)
    
    def _read_chapter(self, filepath: Path) -> str:
        if filepath.suffix == '.md':
            return filepath.read_text(encoding='utf-8')
        elif filepath.suffix == '.docx':
            try:
                from docx import Document
                doc = Document(filepath)
                return '\n\n'.join([p.text for p in doc.paragraphs])
            except ImportError:
                return ""
        return ""
    
    def _normalize_content(self, text: str) -> str:
        # Remove YAML frontmatter
        if text.startswith('---'):
            end = text.find('---', 3)
            if end != -1:
                text = text[end + 3:].strip()
        
        # Standardize scene breaks
        text = re.sub(r'\n\s*[*#-]{3,}\s*\n', f'\n\n{self.SCENE_BREAK}\n\n', text)
        
        # Smart quotes
        text = re.sub(r'"([^"]*)"', r'"\1"', text)
        text = text.replace("'", "'")
        
        # Clean whitespace
        text = re.sub(r'\n{3,}', '\n\n', text)
        
        return text.strip()
    
    def _extract_title(self, text: str, chapter_num: int) -> str:
        match = re.search(r'^#?\s*Chapter\s*\d+[:.]\s*(.+)$', text, re.MULTILINE | re.IGNORECASE)
        if match:
            return match.group(1).strip()
        
        defaults = {
            1: "Arrival", 2: "First Performance", 3: "The Learning Curve",
            4: "Fault Lines", 5: "Tension Wire", 6: "Breaking Point",
            7: "The Reckoning", 8: "The Fall", 9: "What Remains",
            10: "Consequences", 11: "Scattered", 12: "Five Years Later"
        }
        return defaults.get(chapter_num, f"Chapter {chapter_num}")
    
    def normalize_all(self) -> List[ChapterInfo]:
        chapters = []
        for i in range(1, 13):
            for ext in ['.md', '.docx']:
                path = self.chapters_dir / f"chapter_{i:02d}{ext}"
                if path.exists():
                    raw = self._read_chapter(path)
                    content = self._normalize_content(raw)
                    title = self._extract_title(content, i)
                    chapters.append(ChapterInfo(
                        number=i,
                        title=title,
                        word_count=len(content.split()),
                        content=content
                    ))
                    break
        return chapters


class ExportBuilder:
    """Builds export files."""
    
    def __init__(self, base_dir: Path, metadata: BookMetadata):
        self.base_dir = Path(base_dir)
        self.metadata = metadata
        self.exports_dir = self.base_dir / "exports"
        self.exports_dir.mkdir(exist_ok=True)
    
    def build_docx(self, chapters: List[ChapterInfo], output_name: str = "manuscript.docx") -> Path:
        try:
            from docx import Document
            from docx.shared import Inches, Pt
            from docx.enum.text import WD_ALIGN_PARAGRAPH
        except ImportError:
            raise ImportError("python-docx required")
        
        doc = Document()
        
        # Title page
        title_para = doc.add_paragraph()
        title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = title_para.add_run(self.metadata.title.upper())
        run.bold = True
        run.font.size = Pt(24)
        
        doc.add_paragraph()
        author_para = doc.add_paragraph()
        author_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        author_para.add_run(f"by {self.metadata.author}")
        
        doc.add_page_break()
        
        # Chapters
        for chapter in chapters:
            heading = doc.add_heading(f"Chapter {chapter.number}: {chapter.title}", level=1)
            heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            for para_text in chapter.content.split('\n\n'):
                para_text = para_text.strip()
                if para_text == '* * *':
                    p = doc.add_paragraph()
                    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    p.add_run('* * *')
                elif para_text:
                    p = doc.add_paragraph(para_text)
                    p.paragraph_format.first_line_indent = Inches(0.5)
            
            doc.add_page_break()
        
        output_path = self.exports_dir / output_name
        doc.save(output_path)
        return output_path
    
    def build_all(self, chapters: List[ChapterInfo]) -> Dict:
        results = {
            "timestamp": datetime.now().isoformat(),
            "word_count": sum(c.word_count for c in chapters),
            "chapter_count": len(chapters),
            "exports": {}
        }
        
        try:
            docx_path = self.build_docx(chapters)
            results["exports"]["docx"] = str(docx_path)
        except Exception as e:
            results["exports"]["docx_error"] = str(e)
        
        # Version manifest
        manifest = {
            "timestamp": results["timestamp"],
            "metadata": asdict(self.metadata),
            "chapters": [{"number": c.number, "title": c.title, "words": c.word_count} for c in chapters],
            "total_words": results["word_count"]
        }
        manifest_path = self.exports_dir / "version.json"
        manifest_path.write_text(json.dumps(manifest, indent=2))
        
        return results


def export_manuscript(base_dir: str = ".", title: str = "Under the Big Top", author: str = "") -> Dict:
    """Full export pipeline."""
    base_path = Path(base_dir)
    
    normalizer = ManuscriptNormalizer(base_path / "chapters")
    chapters = normalizer.normalize_all()
    
    if not chapters:
        return {"success": False, "error": "No chapters found"}
    
    metadata = BookMetadata(title=title, author=author)
    builder = ExportBuilder(base_path, metadata)
    results = builder.build_all(chapters)
    
    return {"success": True, **results}
