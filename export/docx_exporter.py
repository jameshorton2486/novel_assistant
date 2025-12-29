"""
DOCX Exporter - Standard manuscript format for agents.

Format: Times New Roman 12pt, double-spaced, 1" margins
Industry standard for agent/editor submissions.
"""

from pathlib import Path
from typing import List, Dict, Optional
from datetime import datetime

try:
    from docx import Document
    from docx.shared import Pt, Inches, Twips
    from docx.enum.text import WD_LINE_SPACING, WD_ALIGN_PARAGRAPH
    from docx.enum.style import WD_STYLE_TYPE
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False


class DocxExporter:
    """
    Exports manuscript in standard submission format.
    
    Standard Format:
    - Font: Times New Roman, 12pt
    - Spacing: Double-spaced
    - Margins: 1 inch all around
    - Header: Author name / Title / Page number
    - First line indent: 0.5 inches
    - Scene breaks: Centered # symbol
    
    Usage:
        exporter = DocxExporter()
        
        # Export single chapter
        exporter.export_chapter(text, "Chapter 1", "chapter1.docx")
        
        # Export full manuscript
        exporter.export_full_manuscript(chapters, metadata, "manuscript.docx")
    """
    
    def __init__(self):
        if not DOCX_AVAILABLE:
            raise ImportError(
                "python-docx is required for DOCX export. "
                "Install with: pip install python-docx"
            )
        
        self.font_name = "Times New Roman"
        self.font_size = Pt(12)
        self.margins = Inches(1)
        self.first_line_indent = Inches(0.5)
    
    def _setup_document(self, doc: Document, title: str = None, 
                        author: str = None) -> Document:
        """Apply standard manuscript formatting to document."""
        # Set margins
        for section in doc.sections:
            section.top_margin = self.margins
            section.bottom_margin = self.margins
            section.left_margin = self.margins
            section.right_margin = self.margins
            
            # Add header with page numbers
            header = section.header
            header_para = header.paragraphs[0] if header.paragraphs else header.add_paragraph()
            
            if author and title:
                header_para.text = f"{author} / {title} / "
            
            # Add page number field (simplified)
            header_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        
        return doc
    
    def _format_paragraph(self, para, is_first: bool = False):
        """Apply standard formatting to a paragraph."""
        # Set font
        for run in para.runs:
            run.font.name = self.font_name
            run.font.size = self.font_size
        
        # Set paragraph formatting
        para.paragraph_format.line_spacing_rule = WD_LINE_SPACING.DOUBLE
        para.paragraph_format.space_after = Pt(0)
        para.paragraph_format.space_before = Pt(0)
        
        # First line indent (except for first paragraph of chapter/section)
        if not is_first:
            para.paragraph_format.first_line_indent = self.first_line_indent
    
    def export_chapter(self, chapter_text: str, chapter_name: str,
                       output_path: str, author: str = None) -> str:
        """
        Export single chapter as DOCX.
        
        Args:
            chapter_text: The chapter content
            chapter_name: Name/title of the chapter
            output_path: Where to save the file
            author: Optional author name for header
            
        Returns:
            Path to the exported file
        """
        doc = Document()
        doc = self._setup_document(doc, chapter_name, author)
        
        # Add chapter title
        title = doc.add_heading(chapter_name, level=1)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in title.runs:
            run.font.name = self.font_name
            run.font.size = Pt(14)
            run.font.bold = True
        
        # Add blank line after title
        doc.add_paragraph()
        
        # Process paragraphs
        paragraphs = chapter_text.split('\n\n')
        
        for i, para_text in enumerate(paragraphs):
            para_text = para_text.strip()
            if not para_text:
                continue
            
            # Check for scene break
            if para_text in ['#', '* * *', '***', '---']:
                p = doc.add_paragraph('#')
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                continue
            
            # Add paragraph
            p = doc.add_paragraph(para_text)
            self._format_paragraph(p, is_first=(i == 0))
        
        # Save
        doc.save(output_path)
        return output_path
    
    def export_full_manuscript(self, chapters: List[Dict], 
                                metadata: Dict,
                                output_path: str) -> str:
        """
        Export complete manuscript.
        
        Args:
            chapters: List of {"name": str, "content": str}
            metadata: {"title": str, "author": str, "contact": str, "word_count": int}
            output_path: Where to save the file
            
        Returns:
            Path to the exported file
        """
        doc = Document()
        
        title = metadata.get("title", "Untitled")
        author = metadata.get("author", "")
        
        doc = self._setup_document(doc, title, author)
        
        # Title page
        self._add_title_page(doc, metadata)
        
        # Add chapters
        for i, chapter in enumerate(chapters):
            # Page break before each chapter (except after title page)
            doc.add_page_break()
            
            # Chapter heading
            heading = doc.add_heading(chapter["name"], level=1)
            heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in heading.runs:
                run.font.name = self.font_name
                run.font.size = Pt(14)
            
            # Blank line
            doc.add_paragraph()
            
            # Chapter content
            paragraphs = chapter["content"].split('\n\n')
            
            for j, para_text in enumerate(paragraphs):
                para_text = para_text.strip()
                if not para_text:
                    continue
                
                # Scene break
                if para_text in ['#', '* * *', '***', '---']:
                    p = doc.add_paragraph('#')
                    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    continue
                
                # Regular paragraph
                p = doc.add_paragraph(para_text)
                self._format_paragraph(p, is_first=(j == 0))
        
        # Add "THE END"
        doc.add_paragraph()
        end_para = doc.add_paragraph("THE END")
        end_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in end_para.runs:
            run.font.name = self.font_name
            run.font.size = self.font_size
        
        doc.save(output_path)
        return output_path
    
    def _add_title_page(self, doc: Document, metadata: Dict):
        """Add standard title page."""
        # Top section - contact info (left aligned)
        contact = metadata.get("contact", "")
        if contact:
            p = doc.add_paragraph(contact)
            p.paragraph_format.space_after = Pt(0)
        
        # Word count (top right - approximated with tabs)
        word_count = metadata.get("word_count", 0)
        if word_count:
            p = doc.add_paragraph()
            p.add_run(f"\t\t\t\t\t\tApprox. {word_count:,} words")
            p.paragraph_format.space_after = Pt(0)
        
        # Vertical space
        for _ in range(12):
            doc.add_paragraph()
        
        # Title (centered, larger)
        title = metadata.get("title", "UNTITLED")
        p = doc.add_paragraph(title.upper())
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in p.runs:
            run.font.name = self.font_name
            run.font.size = Pt(14)
            run.font.bold = True
        
        # "by" line
        p = doc.add_paragraph("by")
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in p.runs:
            run.font.name = self.font_name
            run.font.size = self.font_size
        
        # Author name
        author = metadata.get("author", "")
        p = doc.add_paragraph(author)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in p.runs:
            run.font.name = self.font_name
            run.font.size = self.font_size
    
    def calculate_word_count(self, chapters: List[Dict]) -> int:
        """Calculate total word count for manuscript."""
        total = 0
        for chapter in chapters:
            total += len(chapter.get("content", "").split())
        return total
    
    def validate_format(self, file_path: str) -> Dict:
        """
        Validate that a DOCX file meets submission standards.
        
        Args:
            file_path: Path to the DOCX file
            
        Returns:
            Validation results with any issues
        """
        doc = Document(file_path)
        
        issues = []
        
        # Check margins
        for section in doc.sections:
            if abs(section.left_margin.inches - 1.0) > 0.1:
                issues.append("Left margin should be 1 inch")
            if abs(section.right_margin.inches - 1.0) > 0.1:
                issues.append("Right margin should be 1 inch")
        
        # Check fonts in first few paragraphs
        for i, para in enumerate(doc.paragraphs[:10]):
            for run in para.runs:
                if run.font.name and run.font.name != "Times New Roman":
                    issues.append(f"Non-standard font found: {run.font.name}")
                    break
        
        return {
            "valid": len(issues) == 0,
            "issues": issues,
            "paragraph_count": len(doc.paragraphs)
        }
